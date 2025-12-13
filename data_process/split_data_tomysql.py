"""法律法规与文书切分、入库脚本。

整体流程：
1) 遍历 Laws 目录中的 .doc / .docx，按编/章/节/条切分为 Chunk。
2) 将切分结果写入 MySQL（law_documents、law_chunks）。
3) 解析 writ.md，按文书拆分后写入 writ 表。

- 模块 1: 法律文档切分（正则、解析、Chunk 组装）。
- 模块 2: SQLModel ORM 定义与数据库管理工具。
- 模块 3: 文书解析（目录、正文拆分）。
- 模块 4: 数据入库主流程（import_writ_md、main）。
"""

import glob
import os
import re
import json
import logging
from datetime import datetime
from dataclasses import dataclass
from typing import List, Optional

import docx
from sqlmodel import Field, Session, SQLModel, create_engine
from sqlalchemy import Column, Text

from config import settings
from data_process.convert_docx import convert_doc_to_docx

# 配置简单的日志打印，方便调试
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 模块 1: 法律文档切分逻辑

# --- 正则表达式定义 ---
# 针对法条层级（编/章/节/条）定义匹配模式，便于快速识别结构
REGEX_PART = re.compile(r"^第[零一二三四五六七八九十百]+编")  # 编
REGEX_CHAPTER = re.compile(r"^第[零一二三四五六七八九十百]+章")  # 章
REGEX_SECTION = re.compile(r"^第[零一二三四五六七八九十百]+节")  # 节
REGEX_ARTICLE = re.compile(r"^第[零一二三四五六七八九十百千\d]+条")  # 条


def process_law_file(file_path):
    """读取法律文件并切分为 Chunk 列表。

    步骤概览：
    1) 如遇 .doc 先转为临时 .docx，再用 python-docx 读取。
    2) 逐段遍历，依次识别“编/章/节/条”层级；每遇到新条即开始新的 Chunk。
    3) 组合路径（编章节条）+ 正文生成 embedding_text，方便后续向量化。
    4) 返回标准化的 chunk 字典列表；转换产生的临时文件会被清理。
    """

    actual_file_to_read = file_path  # 真正用于读取的文件路径，可能是临时 .docx
    is_temp_file = False             # 标记是否是转换生成的临时文件

    # --- 1. 格式检查与转换 ---
    file_ext = os.path.splitext(file_path)[1].lower()  # 统一小写后缀，便于判断

    if file_ext == '.doc':
        logging.info(f"检测到旧版格式 (.doc)，尝试转换...")
        converted_path = convert_doc_to_docx(file_path)
        if converted_path and os.path.exists(converted_path):
            actual_file_to_read = converted_path
            is_temp_file = True  # 标记为临时文件，处理完后可选择删除
        else:
            logging.error(f"跳过文件 {file_path}: 格式转换失败")
            return []  # 转换失败时直接返回空列表

    # --- 2. 使用 python-docx 读取文件对象 ---
    try:
        doc = docx.Document(actual_file_to_read)
    except Exception as e:
        # 读取失败（如损坏、权限问题）直接跳过
        logging.error(f"无法读取文件 {actual_file_to_read}: {e}")
        return []

    # --- 3. 从原始文件名提取基础信息 ---
    # 注意：我们要用原始 file_path 的文件名，而不是临时文件的
    filename = os.path.basename(file_path)  # 使用原始文件名做元数据

    # 使用 os.path.splitext 安全去除后缀 (.doc 或 .docx)
    # "民法典_2020.docx" -> ("民法典_2020", ".docx")
    filename_no_ext = os.path.splitext(filename)[0]
    law_title = filename_no_ext.split("_")[0]  # 约定文件名以“法名_其他”形式

    chunks = []

    # --- 状态机 (State Machine) 初始化 ---
    # 记录当前层级状态，匹配到更高层级时会重置下级
    state = {
        "part": "",   # 第几编
        "chapter": "",  # 第几章
        "section": ""  # 第几节
    }

    # 暂存当前条文内容，遇到下一条或文件结束时写入列表
    current_article = {
        "id": "",
        "content": ""
    }

    # --- 4. 逐行遍历 Word 文档的段落 ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue  # 跳过空行

        clean_text = re.sub(r'\s+', '', text)  # 去除空白后便于正则匹配

        # A. 层级识别
        # 匹配第几编
        if REGEX_PART.match(clean_text):
            state["part"] = text
            state["chapter"] = ""
            state["section"] = ""
            continue

        # 匹配第几章
        if REGEX_CHAPTER.match(clean_text):
            state["chapter"] = text
            state["section"] = ""
            continue

        # 匹配第几节
        if REGEX_SECTION.match(clean_text):
            state["section"] = text
            continue

        # B. 法条切分
        if REGEX_ARTICLE.match(clean_text):
            # 进入新条文前，把上一条写入列表
            if current_article["id"]:
                save_chunk_to_list(chunks, law_title, state, current_article)

            parts = text.split(maxsplit=1)   # 按空格分割文本，最多分割 1 次
            current_article["id"] = parts[0]
            current_article["content"] = parts[1] if len(parts) > 1 else ""

        # C. 正文累积
        else:
            # 同一条文的正文追加到 content
            if current_article["id"]:
                current_article["content"] += "\n" + text

    # 循环结束收尾
    if current_article["id"]:
        save_chunk_to_list(chunks, law_title, state, current_article)

    # --- 5. 清理临时文件 (可选) ---
    # 如果是你生成的 .docx 临时文件，建议处理完后删除，保持文件夹整洁
    if is_temp_file and os.path.exists(actual_file_to_read):
        try:
            os.remove(actual_file_to_read)
            logging.info(f"清理临时文件: {actual_file_to_read}")
        except OSError:
            pass

    return chunks


def save_chunk_to_list(chunks_list, law_title, state, article_data):
    """将当前条文写入 chunks_list，形成统一的 Chunk 结构。"""
    path_components = [
        law_title,
        state["part"],
        state["chapter"],
        state["section"],
        article_data["id"]
    ]
    valid_paths = [p for p in path_components if p]  # 过滤空层级，保持路径简洁
    path_str = " ".join(valid_paths)
    embedding_text = f"{path_str} ： {article_data['content']}"  # 便于向量化的上下文串

    chunks_list.append({
        "part": state["part"],
        "chapter": state["chapter"],
        "section": state["section"],
        "article_id": article_data["id"],
        "content_text": article_data["content"],
        "embedding_text": embedding_text
    })


# 模块 2: 使用 SQLModel 定义表模型与数据库管理类（法律法规 + 文书）

class LawDocument(SQLModel, table=True):
    """对应表 law_documents 的 ORM 模型。"""

    __tablename__ = "law_documents"

    id: int | None = Field(default=None, primary_key=True)
    filename: str
    law_title: str
    file_path: str
    upload_date: datetime = Field(default_factory=datetime.now)


class LawChunk(SQLModel, table=True):
    """对应表 law_chunks 的 ORM 模型。"""

    __tablename__ = "law_chunks"

    id: int | None = Field(default=None, primary_key=True)
    document_id: int
    law_filename: str
    law_title: str
    part_name: str | None = None   # 编
    chapter_name: str | None = None  # 章
    section_name: str | None = None  # 节
    article_id: str                  # 条
    content_text: str
    embedding_text: str

class WritRecord(SQLModel, table=True):
    """对应表 writ，用于存放文书拆分后的记录。"""

    __tablename__ = "writ"

    id: Optional[int] = Field(default=None, primary_key=True)
    catalogue: str
    title: str
    indexbytitle: str
    context: str = Field(sa_column=Column(Text))


def build_engine(db_config):
    """根据配置构造 SQLModel Engine（MySQL）。"""
    db_url = (
        f"mysql+pymysql://{db_config['user']}:{db_config['password']}"
        f"@{db_config['host']}:{db_config['port']}/{db_config['database']}?charset=utf8mb4"
    )
    # pool_recycle 防止长连接超时；echo=False 关闭 SQL 打印
    return create_engine(db_url, echo=False, pool_recycle=3600)


class LawSQLModelManager:
    """封装基于 SQLModel 的会话创建与批量写入。"""

    def __init__(self, db_config):
        self.db_config = db_config
        self.engine = None

    def connect(self):
        """创建 Engine 并自动建表（包含 law_* 与 writ）。"""
        self.engine = build_engine(self.db_config)
        SQLModel.metadata.create_all(self.engine)
        logging.info("SQLModel 引擎创建成功")

    def close(self):
        """释放底层连接资源。"""
        if self.engine:
            try:
                self.engine.dispose()
            except Exception:
                pass
        self.engine = None

    def insert_document(self, filename, law_title, file_path):
        """插入法律文件记录并返回自增主键。"""
        with Session(self.engine) as session:
            doc = LawDocument(
                filename=filename,
                law_title=law_title,
                file_path=file_path,
            )
            session.add(doc)
            session.commit()
            session.refresh(doc)
            return doc.id

    def insert_chunks_batch(self, chunks_data):
        """批量插入 chunk，向量序列化为 JSON 字符串后存入 TEXT。"""
        objects = []
        for c in chunks_data:
            raw_vector = c.get('embedding', [])
            vector_json_str = json.dumps(raw_vector)
            objects.append(LawChunk(
                document_id=c['document_id'],
                law_filename=c['law_filename'],
                law_title=c['law_title'],
                part_name=c['part'],
                chapter_name=c['chapter'],
                section_name=c['section'],
                article_id=c['article_id'],
                content_text=c['content_text'],
                embedding_text=c['embedding_text'],
                embedding_vector=vector_json_str
            ))

        with Session(self.engine) as session:
            session.add_all(objects)
            session.commit()
            logging.info(f"成功批量插入 {len(objects)} 条数据")


# 模块 3:
# 文书 (writ.md) 解析与入库
# ----------------------------
# 文书目录/正文解析工具
# 以 "|...|Col2|" 的表格标题行标记文书开头，
# 以 "#### xxxx人" 的签字/盖章日期行标记文书结尾，
# 标题后的首个页码行用于定位所属目录。
# 表格行会去重、清理空列并保留分隔行。
# ----------------------------


@dataclass
class CatalogueEntry:
    page: int
    title: str


def parse_catalogue(lines: List[str]) -> List[CatalogueEntry]:
    """解析 writ.md 目录区，提取标题与起始页码。

    规则：
    - 目录位于正文之前，遇到首个表格行即可停止解析；
    - 每行形如 “某某……123”，用正则提取标题与页码；
    - 返回按页码升序的列表，供正文页码反推所属目录。
    """
    entries: List[CatalogueEntry] = []
    pattern = re.compile(r"^(?P<title>.+?)[\.·。……\s]*?(?P<page>\d+)\s*$")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("|"):
            break
        m = pattern.match(stripped)
        if m:
            try:
                page = int(m.group("page"))
            except ValueError:
                continue
            title = m.group("title").strip()
            entries.append(CatalogueEntry(page=page, title=title))
    entries.sort(key=lambda x: x.page)
    return entries


def locate_catalogue(page: Optional[int], catalogues: List[CatalogueEntry]) -> str:
    """
    根据页码找到所属目录标题。

    逻辑：在目录列表中找到最后一个 page <= 当前页的条目；若不存在返回 "未知目录"。
    用途：正文中仅有页码数字行，通过最近页码反推所属目录。
    """
    if page is None or not catalogues:
        return "未知目录"
    best = None
    for entry in catalogues:
        if entry.page <= page:
            best = entry
        else:
            break
    return best.title if best else "未知目录"


def split_writs(lines: List[str], catalogues: List[CatalogueEntry]):
    """切分 writ.md，生成文书记录列表。

    关键规则：
    - 标题行：形如 "|xxx|Col2|"，仅在未进入文书时视为标题；进入正文后同样格式视为正文行；
    - 页码行：纯数字行，标题出现后的首个页码用于定位 catalogue；
    - 结束行：匹配 ^####\s*\S{1,6}人，先写入正文，再结束当前文书；
    - 表格行：去重重复列，清理空列，分隔行保持原样；
    - 空行：直接跳过；所有正文会在 flush 时 strip 前后空白。
    """
    results = []
    current_page: Optional[int] = None
    current_chunk_lines: List[str] = []
    current_title: Optional[str] = None
    current_catalogue: Optional[str] = None  # 标题出现后的首个页码决定

    title_pattern = re.compile(r"^\|\s*(?P<title>[^|]+?)\s*\|\s*Col2\b.*\|", re.IGNORECASE)
    end_pattern = re.compile(r"^####\s*\S{1,6}人", re.IGNORECASE)

    def normalize_table_line(raw_line: str) -> str:
        """
        表格行去重并清理空列，保留分隔行。

        处理步骤：
        1) 非表格行直接返回原行；
        2) 分隔行 (|---|---|) 原样返回；
        3) 普通表格行：
           - 去除重复列（内容相同视为重复，仅保留首个）；
           - 删除空列；
           - 若全空则返回空串，表示丢弃该行。
        """
        stripped = raw_line.strip()
        if not stripped.startswith("|"):
            return raw_line

        parts = raw_line.rstrip("\n").split("|")
        leading = parts[0] == ""
        trailing = parts[-1] == ""
        cells = parts[1:-1] if leading and trailing else parts[1:] if leading else parts[:-1] if trailing else parts

        if not cells:
            return raw_line.rstrip("\n")

        if all(set(c.strip()) == {'-'} for c in cells if c.strip()):
            return "|" + "|".join(cells) + "|"

        seen = set()
        deduped = []
        for cell in cells:
            key = cell.strip()
            if key in seen:
                continue
            seen.add(key)
            deduped.append(cell)

        deduped = [c for c in deduped if c.strip()]
        if not deduped:
            return ""

        return "|" + "|".join(deduped) + "|"

    def flush_chunk():
        """
        将当前文书缓冲写入 results。

        仅在已识别标题且正文非空时写出；context 保留原始格式并去除首尾空白；
        indexbytitle = catalogue-title 便于唯一索引。
        """
        if not current_title or not current_chunk_lines:
            return
        context = "\n".join(current_chunk_lines).strip()
        results.append({
            "catalogue": current_catalogue,
            "title": current_title,
            "indexbytitle": f"{current_catalogue}-{current_title}",
            "context": context
        })

    for raw in lines:
        line = raw.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.isdigit():
            try:
                current_page = int(stripped)
                if current_title and current_catalogue is None:
                    current_catalogue = locate_catalogue(current_page, catalogues)
            except ValueError:
                pass
            continue

        if end_pattern.match(stripped):
            if current_title:
                current_chunk_lines.append(line)
                flush_chunk()
            current_title = None
            current_chunk_lines = []
            current_catalogue = None
            continue

        m = title_pattern.match(stripped)
        if m:
            if not current_title:
                current_title = m.group("title").strip()
                current_chunk_lines = [normalize_table_line(line)]
                current_catalogue = None
                continue
            normalized = normalize_table_line(line)
            current_chunk_lines.append(normalized)
            continue

        if current_title:
            normalized = normalize_table_line(line)
            if normalized:
                current_chunk_lines.append(normalized)

    flush_chunk()
    return results


# ----------------------------
# 文书入库
# ----------------------------

def import_writ_md(db_manager, md_path):
    """读取 writ.md，切分文书并写入 writ 表。

    流程：加载文件 -> 解析目录 -> 切分正文 -> 批量写入数据库。
    遇到缺失文件会直接返回并输出错误日志。
    """
    if not os.path.exists(md_path):
        logging.error(f"未找到 markdown 文件: {md_path}")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    catalogues = parse_catalogue(lines)
    chunks = split_writs(lines, catalogues)
    logging.info(f"解析得到 {len(chunks)} 篇文书")

    if not chunks:
        logging.warning("未切分出任何文书，程序结束")
        return

    with Session(db_manager.engine) as session:
        objects = [WritRecord(**c) for c in chunks]
        session.add_all(objects)
        session.commit()
        logging.info(f"成功写入 {len(objects)} 条记录到表 writ")
# ==============================================================================

def main():
    # 1. 数据库连接配置
    DB_CONFIG = {
        "host": settings.HOST,
        "user": settings.USENAME,
        "password": settings.PASSWORD,
        "database": settings.DATABASE,
        "port": settings.PORT
    }

    # 2. 扫描法律文件
    files = glob.glob(os.path.join("Laws", "*.doc*"))

    db_manager = LawSQLModelManager(DB_CONFIG)

    try:
        db_manager.connect()

        # --- 先处理法律法规 ---
        for file_path in files:
            if os.path.basename(file_path).startswith("~$"):
                continue
            if not os.path.exists(file_path):
                continue

            logging.info(f"=== 开始处理法律文件: {file_path} ===")

            filename = os.path.basename(file_path)
            law_title = os.path.splitext(filename)[0].split("_")[0]

            chunks = process_law_file(file_path)
            if not chunks:
                logging.warning("未提取到有效内容或格式转换失败，跳过")
                continue

            doc_id = db_manager.insert_document(filename, law_title, file_path)
            logging.info(f"文档记录创建成功，ID: {doc_id}")

            for chunk in chunks:
                chunk['document_id'] = doc_id
                chunk['law_filename'] = filename
                chunk['law_title'] = law_title
            db_manager.insert_chunks_batch(chunks)

        # --- 再处理文书 ---
        md_path = os.path.join(os.path.dirname(__file__), "writ", "writ.md")
        logging.info("=== 开始处理 writ.md 文书 ===")
        import_writ_md(db_manager, md_path)

    except Exception as e:
        logging.error(f"程序运行发生致命错误: {e}")
    finally:
        db_manager.close()

if __name__ == "__main__":
    main()